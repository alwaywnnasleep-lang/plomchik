import io
from datetime import datetime
from django.utils import timezone
from rest_framework.views import APIView
from rest_framework.response import Response
from rest_framework import status
from django.http import FileResponse
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

from apps.tasks.models import Task
from apps.structure.services import get_units_under_authority
from apps.users.permissions import IsHeadOrAbove


class ReportGenerateView(APIView):
    permission_classes = [IsHeadOrAbove]

    def get(self, request):
        # Параметры фильтрации
        start_date = request.query_params.get('start_date')
        end_date = request.query_params.get('end_date')
        unit_id = request.query_params.get('unit_id')
        status_filter = request.query_params.get('status')

        user = request.user
        unit_ids = get_units_under_authority(user)

        qs = Task.objects.select_related('org_unit', 'assigned_to', 'created_by')
        if user.role not in ('commander', 'deputy_commander'):
            qs = qs.filter(org_unit_id__in=unit_ids)

        if start_date:
            qs = qs.filter(created_at__gte=start_date)
        if end_date:
            qs = qs.filter(created_at__lte=end_date)
        if unit_id:
            qs = qs.filter(org_unit_id=unit_id)
        if status_filter:
            qs = qs.filter(status=status_filter)

        total = qs.count()
        completed = qs.filter(status='done').count()
        in_progress = qs.filter(status='in_progress').count()
        overdue = qs.filter(deadline__lt=timezone.now()).exclude(status='done').count()

        priorities = {
            'Критический': qs.filter(priority='critical').count(),
            'Высокий': qs.filter(priority='high').count(),
            'Средний': qs.filter(priority='medium').count(),
            'Низкий': qs.filter(priority='low').count(),
        }

        tasks = qs.order_by('-created_at')[:200]

        # Создаём документ
        doc = Document()

        # Заголовок
        title = doc.add_heading('Отчёт по задачам', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Период
        period_str = ''
        if start_date and end_date:
            period_str = f'с {start_date} по {end_date}'
        elif start_date:
            period_str = f'с {start_date}'
        elif end_date:
            period_str = f'до {end_date}'
        else:
            period_str = 'за всё время'
        doc.add_paragraph(f'Период: {period_str}')
        doc.add_paragraph(f'Дата формирования: {datetime.now().strftime("%d.%m.%Y %H:%M")}')
        doc.add_paragraph(f'Пользователь: {user.full_name} ({user.rank})')
        doc.add_paragraph()

        # Сводная статистика
        doc.add_heading('Общая статистика', level=1)
        stats_table = doc.add_table(rows=1, cols=2)
        stats_table.style = 'Light Shading'
        hdr = stats_table.rows[0].cells
        hdr[0].text = 'Показатель'
        hdr[1].text = 'Значение'

        for label, value in [
            ('Всего задач', total),
            ('Выполнено', completed),
            ('В работе', in_progress),
            ('Просрочено', overdue),
        ]:
            row = stats_table.add_row().cells
            row[0].text = label
            row[1].text = str(value)

        doc.add_paragraph()

        # Приоритеты
        doc.add_heading('Распределение по приоритетам', level=1)
        prio_table = doc.add_table(rows=1, cols=2)
        prio_table.style = 'Light Shading'
        prio_hdr = prio_table.rows[0].cells
        prio_hdr[0].text = 'Приоритет'
        prio_hdr[1].text = 'Количество'

        for prio, cnt in priorities.items():
            row = prio_table.add_row().cells
            row[0].text = prio
            row[1].text = str(cnt)

        doc.add_paragraph()

        # Список задач
        doc.add_heading('Список задач', level=1)
        task_table = doc.add_table(rows=1, cols=6)
        task_table.style = 'Light Shading'
        task_hdr = task_table.rows[0].cells
        task_hdr[0].text = 'ID'
        task_hdr[1].text = 'Название'
        task_hdr[2].text = 'Статус'
        task_hdr[3].text = 'Приоритет'
        task_hdr[4].text = 'Дедлайн'
        task_hdr[5].text = 'Подразделение'

        for task in tasks:
            row = task_table.add_row().cells
            row[0].text = str(task.id)
            row[1].text = task.title
            row[2].text = task.get_status_display()
            row[3].text = task.get_priority_display()
            row[4].text = task.deadline.strftime('%d.%m.%Y %H:%M') if task.deadline else '—'
            row[5].text = task.org_unit.name if task.org_unit else '—'

        # Сохраняем в поток
        file_stream = io.BytesIO()
        doc.save(file_stream)
        file_stream.seek(0)

        filename = f'report_{datetime.now().strftime("%Y%m%d_%H%M")}.docx'
        response = FileResponse(file_stream, as_attachment=True, filename=filename)
        response['Content-Type'] = 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        return response