# backend/apps/autoplan/parsers.py

import re
import logging
from datetime import datetime, timedelta
from typing import List, Dict, Any

logger = logging.getLogger(__name__)


# ----------------------------------------------------------------------
# Вспомогательные функции для работы с датами (из military_parser.py)
# ----------------------------------------------------------------------

def normalize_date_str(s: str) -> str:
    """Преобразует '7.10---6.11' или '22--24' в строку диапазона."""
    s = re.sub(r'[—–\-]+', '-', s).strip()
    if '-' in s:
        parts = s.split('-')
        if len(parts) == 2:
            start = parts[0].strip()
            end = parts[1].strip()
            return f"{start}|{end}"
    return s


def parse_single_date(expr: str, base_year: int, base_month: int):
    """Парсит '7.10' или '7' в datetime."""
    if '.' in expr:
        parts = expr.split('.')
        day = int(parts[0])
        month = int(parts[1]) if len(parts) > 1 else base_month
        year = base_year
    else:
        day = int(expr)
        month = base_month
        year = base_year
    try:
        return datetime(year, month, day)
    except ValueError:
        return None


def expand_range(date_expr: str, base_year: int, base_month: int):
    """Возвращает список datetime объектов для диапазона или одиночной даты."""
    if '|' in date_expr:
        start_str, end_str = date_expr.split('|')
        start = parse_single_date(start_str, base_year, base_month)
        end = parse_single_date(end_str, base_year, base_month)
        if start and end:
            delta = (end - start).days
            return [start + timedelta(days=i) for i in range(delta + 1)]
    else:
        d = parse_single_date(date_expr, base_year, base_month)
        return [d] if d else []
    return []


def extract_month_year_from_text(text: str):
    """Ищет 'в октябре 2025 г.' в тексте документа."""
    text_sample = text[:2000]
    match = re.search(r'в\s+(\w+)\s+(\d{4})\s*г\.', text_sample, re.IGNORECASE)
    if match:
        month_name = match.group(1).lower()
        year = int(match.group(2))
        month_map = {
            'января': 1, 'февраля': 2, 'марта': 3, 'апреля': 4,
            'мая': 5, 'июня': 6, 'июля': 7, 'августа': 8,
            'сентября': 9, 'октября': 10, 'ноября': 11, 'декабря': 12
        }
        month = month_map.get(month_name, 10)
        return year, month
    return 2025, 10  # значение по умолчанию


# ----------------------------------------------------------------------
# Загрузка данных из файлов (с нормализацией ячеек)
# ----------------------------------------------------------------------

def parse_docx_table(file_path: str):
    """Извлекает все строки из таблиц DOCX, объединяя многострочный текст в ячейке.

    ВАЖНО: python-docx возвращает дублирующиеся объекты ячеек для объединённых (merged)
    ячеек — одна физическая ячейка появляется столько раз, сколько колонок она занимает.
    Эта функция дедуплицирует ячейки по идентичности XML-элемента (_tc), чтобы
    каждая физическая ячейка учитывалась ровно один раз.
    """
    from docx import Document
    doc = Document(file_path)
    all_rows = []
    full_text = '\n'.join([p.text for p in doc.paragraphs])
    for table in doc.tables:
        for row in table.rows:
            seen_tc = set()
            row_data = []
            for cell in row.cells:
                tc_id = id(cell._tc)
                if tc_id in seen_tc:
                    continue  # пропускаем дубликат объединённой ячейки
                seen_tc.add(tc_id)
                # Заменяем переносы строк на пробелы и удаляем лишние пробелы
                text = ' '.join(cell.text.split())
                row_data.append(text)
            if any(cell for cell in row_data):
                all_rows.append(row_data)
    return all_rows, full_text


def parse_xlsx_data(file_path: str):
    """Извлекает данные из первого листа Excel, объединяя переносы в ячейках."""
    from openpyxl import load_workbook
    wb = load_workbook(file_path, data_only=True)
    ws = wb.active
    data = []
    full_text = ''
    for row in ws.iter_rows(values_only=True):
        row_data = []
        for cell in row:
            if cell is None:
                val = ''
            else:
                if isinstance(cell, (int, float)):
                    val = str(int(cell)) if cell == int(cell) else str(cell)
                else:
                    val = str(cell).strip()
                # Заменяем переносы строк на пробелы
                val = ' '.join(val.split())
            row_data.append(val)
        if any(cell for cell in row_data):
            data.append(row_data)
            full_text += ' '.join(row_data) + '\n'
    wb.close()
    return data, full_text


# ----------------------------------------------------------------------
# Специализированный парсер для военных планов-календарей
# ----------------------------------------------------------------------

def parse_military_calendar_table(file_path: str, file_type: str) -> List[Dict[str, Any]]:
    """
    Парсит план-календарь:
    - находит строку с днями месяца (числа от 1 до 31)
    - первые колонки (до первой колонки с датой) – текстовые: мероприятие, руководитель, состав
    - для каждой строки мероприятия извлекает даты из колонок-дней

    Для DOCX корректно обрабатывает объединённые (merged) ячейки: одна физическая ячейка
    может занимать несколько колонок, и мы используем реальный XML-элемент (_tc) для
    дедупликации — так каждая ячейка обрабатывается ровно один раз, а её «охват» по
    колонкам сопоставляется с картой дней из заголовочной строки.
    """
    if file_type in ('xlsx', 'xls'):
        return _parse_military_calendar_xlsx(file_path)
    elif file_type in ('docx', 'doc'):
        return _parse_military_calendar_docx(file_path)
    else:
        raise ValueError("Неподдерживаемый тип файла для военного календаря")


def _parse_military_calendar_xlsx(file_path: str) -> List[Dict[str, Any]]:
    """Парсинг военного календаря из XLSX (без объединённых ячеек)."""
    rows, full_text = parse_xlsx_data(file_path)
    if not rows:
        logger.warning("Нет данных в таблице XLSX")
        return []

    base_year, base_month = extract_month_year_from_text(full_text)

    # Ищем строку с днями месяца
    date_row_idx = None
    for i, row in enumerate(rows):
        count = sum(1 for cell in row if _is_day(cell))
        if count >= 10:
            date_row_idx = i
            break

    if date_row_idx is None:
        logger.warning("Не найдена строка с днями (XLSX)")
        return []

    date_row = rows[date_row_idx]
    date_columns = [idx for idx, cell in enumerate(date_row) if _is_day(cell)]
    if not date_columns:
        return []

    first_date_col = date_columns[0]
    events = []

    for row in rows[date_row_idx + 1:]:
        if not row or all(cell == '' for cell in row):
            continue
        title = row[0] if len(row) > 0 else ''
        if not title:
            continue
        responsible = row[1] if len(row) > 1 else ''
        personnel = row[2] if len(row) > 2 else ''

        for col_idx in date_columns:
            if col_idx >= len(row):
                continue
            cell_value = row[col_idx]
            if not cell_value:
                continue
            normalized = normalize_date_str(cell_value)
            dates = expand_range(normalized, base_year, base_month)
            if dates:
                events.append(_make_event(title, responsible, personnel, dates))

    return _deduplicate_events(events)


def _parse_military_calendar_docx(file_path: str) -> List[Dict[str, Any]]:
    """
    Парсинг военного календаря из DOCX с корректной обработкой объединённых ячеек.

    Алгоритм:
    1. Строим карту grid-колонка → день из строки-заголовка (с датами).
    2. Для каждой строки данных группируем ячейки по идентичности _tc (XML-элемент),
       определяем какие дни охватывает каждая уникальная ячейка,
       и извлекаем событие по значению ячейки + охваченным дням.
    """
    from docx import Document
    doc = Document(file_path)

    if not doc.tables:
        logger.warning("В документе нет таблиц")
        return []

    # Используем первую таблицу (основная таблица плана-календаря)
    table = doc.tables[0]

    full_text = '\n'.join([p.text for p in doc.paragraphs])
    base_year, base_month = extract_month_year_from_text(full_text)

    # Шаг 1: найти строку с днями месяца и построить карту grid-col -> day
    # Используем raw-итерацию (с дублями) чтобы получить реальные grid-индексы
    date_row_raw_idx = None
    grid_col_to_day: Dict[int, int] = {}

    for row_idx, row in enumerate(table.rows):
        temp_map: Dict[int, int] = {}
        for col_idx, cell in enumerate(row.cells):
            try:
                val = int(cell.text.strip())
                if 1 <= val <= 31 and col_idx not in temp_map:
                    temp_map[col_idx] = val
            except (ValueError, TypeError):
                pass
        if len(temp_map) >= 10:
            date_row_raw_idx = row_idx
            grid_col_to_day = temp_map
            break

    if date_row_raw_idx is None:
        logger.warning("Не найдена строка с днями месяца (DOCX)")
        return []

    day_grid_cols = set(grid_col_to_day.keys())
    first_day_col = min(day_grid_cols)

    # Шаг 2: собираем строки данных с текстовыми полями и датами
    # Структура каждой строки: {'raw_title', 'responsible', 'personnel', 'date_cells'}
    data_rows = []

    for row_idx in range(date_row_raw_idx + 1, len(table.rows)):
        row = table.rows[row_idx]

        # Группируем колонки по уникальному XML-элементу ячейки (_tc)
        tc_to_cols: Dict[Any, List[int]] = {}
        tc_order: List[Any] = []
        for col_idx, cell in enumerate(row.cells):
            tc = cell._tc
            if tc not in tc_to_cols:
                tc_to_cols[tc] = []
                tc_order.append(tc)
            tc_to_cols[tc].append(col_idx)

        # Текстовые поля: ячейки, чьи grid-колонки все левее первой колонки с датой
        text_values = [
            ' '.join(row.cells[cols[0]].text.split())
            for tc, cols in zip(tc_order, [tc_to_cols[t] for t in tc_order])
            if all(c < first_day_col for c in cols)
        ]

        raw_title = text_values[0] if text_values else ''
        if not raw_title:
            continue

        responsible = text_values[1] if len(text_values) > 1 else ''
        personnel = text_values[2] if len(text_values) > 2 else ''

        # Собираем ячейки с датами для этой строки
        date_cells = []
        for tc in tc_order:
            cols = tc_to_cols[tc]
            cell_text = ' '.join(row.cells[cols[0]].text.split())
            if not cell_text:
                continue
            spanned_days = sorted(grid_col_to_day[c] for c in cols if c in day_grid_cols)
            if spanned_days:
                date_cells.append((cell_text, spanned_days))

        data_rows.append({
            'raw_title': raw_title,
            'responsible': responsible,
            'personnel': personnel,
            'date_cells': date_cells,
        })

    # Шаг 3: определяем родителей и дочерние строки, строим итоговые события.
    # Признак дочерней строки: первый символ строчная буква ИЛИ цифра.
    # Родитель, у которого есть хотя бы одна дочерняя строка, сам событий не создаёт —
    # вместо этого каждая дочь создаёт событие с объединённым заголовком «родитель + дочь».
    events = []

    # Сначала помечаем, у каких строк есть дочерние
    has_children = [False] * len(data_rows)
    for i, dr in enumerate(data_rows):
        fc = dr['raw_title'][0] if dr['raw_title'] else ''
        if (fc.islower() or fc.isdigit()) and i > 0:
            # Найти ближайшего родителя (последний не-дочерний выше)
            for j in range(i - 1, -1, -1):
                pfc = data_rows[j]['raw_title'][0] if data_rows[j]['raw_title'] else ''
                if not (pfc.islower() or pfc.isdigit()):
                    has_children[j] = True
                    break

    # Теперь обрабатываем строки
    last_parent: Dict[str, str] = {}  # raw_title, responsible, personnel

    for i, dr in enumerate(data_rows):
        raw_title = dr['raw_title']
        responsible = dr['responsible']
        personnel = dr['personnel']
        date_cells = dr['date_cells']

        fc = raw_title[0] if raw_title else ''
        is_child = bool(last_parent) and (fc.islower() or fc.isdigit())

        if is_child:
            # Объединяем с родительским заголовком
            title = last_parent['raw_title'] + ' ' + raw_title
            if not responsible:
                responsible = last_parent['responsible']
            if not personnel:
                personnel = last_parent['personnel']
        else:
            title = raw_title
            last_parent = {'raw_title': raw_title, 'responsible': responsible, 'personnel': personnel}

            # Если у этого родителя есть дочерние строки — он тоже создаёт события
            # (собственные даты родителя не теряются)
            pass  # продолжаем создавать события ниже

        # Создаём события из заранее собранных ячеек с датами
        for cell_text, spanned_days in date_cells:
            normalized = normalize_date_str(cell_text)
            dates = expand_range(normalized, base_year, base_month)

            if dates:
                events.append(_make_event(title, responsible, personnel, dates))
            else:
                # Текст не парсится как дата: используем охваченные дни напрямую
                start = datetime(base_year, base_month, spanned_days[0])
                end = datetime(base_year, base_month, spanned_days[-1])
                events.append(_make_event(title, responsible, personnel,
                                          [start] if start == end else [start, end]))

    result = _deduplicate_events(events)
    logger.info(f"Military calendar parser (DOCX): найдено {len(result)} мероприятий")
    return result


def _is_day(cell: str) -> bool:
    """Возвращает True, если строка — целое число от 1 до 31."""
    try:
        return 1 <= int(str(cell).strip()) <= 31
    except (ValueError, TypeError):
        return False


def _make_event(title: str, responsible: str, personnel: str,
                dates: List) -> Dict[str, Any]:
    """Формирует словарь события из списка дат."""
    if len(dates) > 1:
        return {
            'title': title,
            'responsible': responsible,
            'personnel': personnel,
            'start_date': dates[0].isoformat(),
            'end_date': dates[-1].isoformat(),
            'is_range': True,
            'date': dates[0].isoformat(),
        }
    return {
        'title': title,
        'responsible': responsible,
        'personnel': personnel,
        'date': dates[0].isoformat(),
        'is_range': False,
    }


def _deduplicate_events(events: List[Dict[str, Any]]) -> List[Dict[str, Any]]:
    """Убирает дублирующиеся события (по названию + дате начала)."""
    unique: Dict[str, Dict[str, Any]] = {}
    for ev in events:
        key = ev['title'] + '|' + ev.get('date', '') + '|' + ev.get('start_date', '')
        if key not in unique:
            unique[key] = ev
    return list(unique.values())


# ----------------------------------------------------------------------
# Общие парсеры для DOCX, XLSX, PDF (базовая реализация)
# ----------------------------------------------------------------------

def parse_docx(file_path: str) -> List[Dict[str, Any]]:
    """Парсинг DOCX-файла: извлекает текст из таблиц и абзацев."""
    from docx import Document

    doc = Document(file_path)
    results = []

    for table in doc.tables:
        if not table.rows:
            continue
        headers = [cell.text.strip().lower() for cell in table.rows[0].cells]

        for row in table.rows[1:]:
            row_data = {}
            for idx, cell in enumerate(row.cells):
                key = headers[idx] if idx < len(headers) else f'col_{idx}'
                row_data[key] = cell.text.strip()

            mapped = _map_row(row_data, headers)
            if mapped and mapped.get('title'):
                results.append(mapped)

    if not results:
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                results.append({
                    'title': text,
                    'deadline': '',
                    'responsible': '',
                    'note': ''
                })
    return results


def parse_xlsx(file_path: str) -> List[Dict[str, Any]]:
    """Парсинг XLSX: извлекает данные из первого листа."""
    from openpyxl import load_workbook

    wb = load_workbook(file_path, data_only=True)
    ws = wb.active
    rows = list(ws.iter_rows(values_only=True))
    if not rows:
        wb.close()
        return []

    headers = [str(h).strip().lower() if h else f'col_{i}' for i, h in enumerate(rows[0])]
    results = []

    for row in rows[1:]:
        row_data = {}
        for idx, val in enumerate(row):
            key = headers[idx] if idx < len(headers) else f'col_{idx}'
            row_data[key] = str(val).strip() if val is not None else ''

        mapped = _map_row(row_data, headers)
        if mapped and mapped.get('title'):
            results.append(mapped)

    wb.close()
    return results


def parse_pdf(file_path: str) -> List[Dict[str, Any]]:
    """Парсинг PDF: извлекает текст со страниц (примитивно)."""
    from PyPDF2 import PdfReader

    reader = PdfReader(file_path)
    full_text = ''
    for page in reader.pages:
        text = page.extract_text()
        if text:
            full_text += text + '\n'

    lines = [l.strip() for l in full_text.split('\n') if l.strip()]
    results = []
    for line in lines:
        if len(line) > 20:
            results.append({
                'title': line,
                'deadline': '',
                'responsible': '',
                'note': ''
            })
    return results


# ----------------------------------------------------------------------
# Вспомогательная функция для маппинга строки в стандартный словарь
# ----------------------------------------------------------------------

HEADER_MAPPING = {
    'number': ['№', '№ п/п', 'номер', 'n', '#'],
    'title': ['мероприятие', 'наименование', 'задача', 'название', 'содержание'],
    'deadline': ['срок', 'дата', 'срок выполнения', 'дедлайн', 'до'],
    'responsible': ['ответственный', 'исполнитель', 'отв.', 'кто'],
    'note': ['примечание', 'прим.', 'замечание', 'комментарий'],
}

def _map_row(row_data: Dict[str, str], headers: List[str]) -> Dict[str, str]:
    """Преобразует строку таблицы в единый формат."""
    mapped = {}
    for field, variants in HEADER_MAPPING.items():
        for header_key, value in row_data.items():
            if any(v in header_key for v in variants):
                mapped[field] = value
                break
    if not mapped.get('title'):
        values = list(row_data.values())
        if values:
            mapped['title'] = values[0]
    return mapped


# ----------------------------------------------------------------------
# Главная функция parse_document
# ----------------------------------------------------------------------

def parse_document(file_path: str, file_type: str) -> List[Dict[str, Any]]:
    """
    Основная точка входа для парсинга документов.
    Сначала пытается распознать военный календарь, иначе использует общий парсер.
    """
    if file_type == 'doc':
        raise ValueError(
            "Файл в старом формате .doc не поддерживается. "
            "Пожалуйста, откройте документ в Microsoft Word, LibreOffice или Google Docs "
            "и сохраните его в формате .docx."
        )

    if file_type in ('docx', 'doc', 'xlsx', 'xls'):
        try:
            military_events = parse_military_calendar_table(file_path, file_type)
            if military_events:
                return military_events
        except Exception as e:
            logger.warning(f"Military calendar parser failed: {e}, falling back to generic parser")

    if file_type == 'docx':
        return parse_docx(file_path)
    elif file_type in ('xlsx', 'xls'):
        return parse_xlsx(file_path)
    elif file_type == 'pdf':
        return parse_pdf(file_path)
    else:
        raise ValueError(f"Неподдерживаемый тип файла: {file_type}")